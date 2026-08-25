"""Build the S3 five-screen report directly as a formatted .docx.

Replicates the visual conventions of test_skill/template.docx:
- Font: Arial (Latin) / 等线 (East Asian), 11pt body
- Document title: 26pt bold
- Screen title (第N屏 · xxx): 18pt bold
- Template section (渠道结构/关键发现/规则诊断/数据缺口 etc.): 16pt bold
- Finding / rule heading (发现：xxx, R61｜xxx): 15pt bold
- Evidence/body bullets: native Word bullet list, blue (#3370ff) bullet glyph
- Tables: single-line borders in #dee0e3, fixed layout, equal column widths,
  cell margins (top 60twips/left 120/bottom 30/right 120 dxa)
- Paragraph spacing: before/after 120 twips, line spacing 288 auto

This module exposes a small structured API (ReportBuilder) so the S3 report
generator can emit .docx directly instead of Markdown, plus a
``markdown_to_docx`` function that parses a Markdown string written by an
agent (headings, bullets, tables, ``---`` page breaks) and renders it through
the same ReportBuilder styling.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

FONT_LATIN = "Arial"
FONT_EA = "等线"
BULLET_COLOR = "3370FF"
BORDER_COLOR = "DEE0E3"
EVIDENCE_PREFIXES = ("【数据事实】", "【规则结论】", "【推断】")

TITLE_SIZE = Pt(26)
SCREEN_SIZE = Pt(18)
SECTION_SIZE = Pt(16)
FINDING_SIZE = Pt(15)
BODY_SIZE = Pt(11)


def _set_east_asian_font(run, name: str = FONT_EA) -> None:
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), FONT_LATIN)
    rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    rFonts.set(qn("w:cs"), FONT_LATIN)


def _spacing(paragraph, before=120, after=120) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    spacing.set(qn("w:line"), "288")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def _enable_bullet(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "1")
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)


class ReportBuilder:
    def __init__(self) -> None:
        self.doc = Document()
        self._setup_styles()
        self._setup_numbering()

    def _setup_styles(self) -> None:
        normal = self.doc.styles["Normal"]
        normal.font.name = FONT_LATIN
        normal.font.size = BODY_SIZE

    def _setup_numbering(self) -> None:
        numbering_part = self.doc.part.numbering_part
        numbering_elm = numbering_part.numbering_definitions._numbering

        abstract_num = OxmlElement("w:abstractNum")
        abstract_num.set(qn("w:abstractNumId"), "1")
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        numFmt = OxmlElement("w:numFmt")
        numFmt.set(qn("w:val"), "bullet")
        lvlText = OxmlElement("w:lvlText")
        lvlText.set(qn("w:val"), "•")
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        pPr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "360")
        ind.set(qn("w:hanging"), "360")
        pPr.append(ind)
        rPr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), BULLET_COLOR)
        rPr.append(color)
        lvl.append(numFmt)
        lvl.append(lvlText)
        lvl.append(suff)
        lvl.append(pPr)
        lvl.append(rPr)
        abstract_num.append(lvl)
        numbering_elm.append(abstract_num)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), "1")
        abstractNumId = OxmlElement("w:abstractNumId")
        abstractNumId.set(qn("w:val"), "1")
        num.append(abstractNumId)
        numbering_elm.append(num)

    def _add_run(self, paragraph, text: str, size: Pt, bold: bool = False, color: RGBColor | None = None):
        run = paragraph.add_run(text)
        run.font.name = FONT_LATIN
        run.font.size = size
        run.bold = bold
        if color:
            run.font.color.rgb = color
        _set_east_asian_font(run)
        return run

    def add_title(self, text: str) -> None:
        p = self.doc.add_paragraph()
        _spacing(p)
        self._add_run(p, text, TITLE_SIZE, bold=True)

    def add_screen_title(self, text: str) -> None:
        p = self.doc.add_paragraph()
        _spacing(p)
        self._add_run(p, text, SCREEN_SIZE, bold=True)

    def add_section(self, text: str) -> None:
        p = self.doc.add_paragraph()
        _spacing(p)
        self._add_run(p, text, SECTION_SIZE, bold=True)

    def add_finding_heading(self, text: str) -> None:
        p = self.doc.add_paragraph()
        _spacing(p)
        self._add_run(p, text, FINDING_SIZE, bold=True)

    def add_body(self, text: str) -> None:
        p = self.doc.add_paragraph()
        _spacing(p)
        prefix = next((p_ for p_ in EVIDENCE_PREFIXES if text.startswith(p_)), None)
        if prefix:
            self._add_run(p, prefix, BODY_SIZE, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
            self._add_run(p, text[len(prefix):], BODY_SIZE, bold=False)
        else:
            self._add_run(p, text, BODY_SIZE, bold=False)

    def add_evidence_label(self, text: str = "证据：") -> None:
        self.add_body(text)

    def add_bullet(self, text: str, emphasis_prefix: str | None = None) -> None:
        p = self.doc.add_paragraph()
        _enable_bullet(p)
        _spacing(p)
        prefix = emphasis_prefix if emphasis_prefix and text.startswith(emphasis_prefix) else None
        if prefix is None:
            prefix = next((p_ for p_ in EVIDENCE_PREFIXES if text.startswith(p_)), None)
        if prefix:
            self._add_run(p, prefix, BODY_SIZE, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
            self._add_run(p, text[len(prefix):], BODY_SIZE, bold=False)
        else:
            self._add_run(p, text, BODY_SIZE, bold=False)

    def add_blank(self) -> None:
        p = self.doc.add_paragraph()
        _spacing(p)

    def add_table(self, headers: list[str], rows: list[list[str]]) -> None:
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        self._style_table(table)
        for col_idx, header in enumerate(headers):
            cell = table.rows[0].cells[col_idx]
            self._fill_cell(cell, header, bold=True)
        for row_idx, row in enumerate(rows, start=1):
            for col_idx, value in enumerate(row):
                cell = table.rows[row_idx].cells[col_idx]
                self._fill_cell(cell, str(value))

    def _style_table(self, table) -> None:
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:color"), BORDER_COLOR)
            el.set(qn("w:sz"), "4")
            borders.append(el)
        tblPr.append(borders)
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tblPr.append(layout)

    def _fill_cell(self, cell, text: str, bold: bool = False) -> None:
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement("w:tcMar")
        for edge, value in (("top", "60"), ("left", "120"), ("bottom", "30"), ("right", "120")):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:type"), "dxa")
            el.set(qn("w:w"), value)
            tcMar.append(el)
        tcPr.append(tcMar)
        cell.text = ""
        p = cell.paragraphs[0]
        _spacing(p)
        self._add_run(p, text, BODY_SIZE, bold=bold)

    def add_page_break(self) -> None:
        self.doc.add_page_break()

    def save(self, path: Path) -> None:
        self.doc.save(path)


_HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)$")
_BULLET_RE = re.compile(r"^[-*]\s+(.*)$")
_TABLE_ROW_RE = re.compile(r"^\|(.+)\|$")
_TABLE_SEP_RE = re.compile(r"^\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)*\|?$")


def _split_table_row(line: str) -> list[str]:
    inner = line.strip()
    if inner.startswith("|"):
        inner = inner[1:]
    if inner.endswith("|"):
        inner = inner[:-1]
    return [cell.strip() for cell in inner.split("|")]


def markdown_to_docx(markdown_text: str, output_path: Path) -> None:
    """Parse a Markdown report written by an agent and render it as .docx.

    Heading levels map onto the ReportBuilder styles reverse-engineered from
    test_skill/template.docx: the first ``#`` becomes the document title,
    every later ``#`` becomes a screen title, ``##`` a template section,
    ``###`` a finding heading. Bullets auto-detect the
    【数据事实】/【规则结论】/【推断】 evidence prefixes. A lone ``---`` line
    inserts a page break (matching how screens are separated in this report).
    GFM tables are rendered with add_table.
    """
    r = ReportBuilder()
    lines = markdown_text.splitlines()
    seen_h1 = False
    i = 0
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            r.add_page_break()
            i += 1
            continue

        heading_match = _HEADING_RE.match(stripped)
        if heading_match:
            hashes, title = heading_match.groups()
            title = title.strip()
            level = len(hashes)
            if level == 1:
                if not seen_h1:
                    r.add_title(title)
                    seen_h1 = True
                else:
                    r.add_screen_title(title)
            elif level == 2:
                r.add_section(title)
            else:
                r.add_finding_heading(title)
            i += 1
            continue

        if _TABLE_ROW_RE.match(stripped) and i + 1 < len(lines) and _TABLE_SEP_RE.match(lines[i + 1].strip()):
            headers = _split_table_row(stripped)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and _TABLE_ROW_RE.match(lines[i].strip()):
                rows.append(_split_table_row(lines[i].strip()))
                i += 1
            r.add_table(headers, rows)
            continue

        bullet_match = _BULLET_RE.match(stripped)
        if bullet_match:
            r.add_bullet(bullet_match.group(1).strip())
            i += 1
            continue

        r.add_body(stripped)
        i += 1

    r.save(output_path)
