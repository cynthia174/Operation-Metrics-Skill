"""Build an evidence-bounded report from current run artifacts."""
import argparse, json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--rule-results', required=True, type=Path)
    ap.add_argument('--preflight', required=True, type=Path)
    ap.add_argument('--output', required=True, type=Path)
    args = ap.parse_args()
    rules = json.loads(args.rule_results.read_text(encoding='utf-8'))
    preflight = json.loads(args.preflight.read_text(encoding='utf-8'))
    discovery = preflight.get('discoveries', [{}])[0]
    date_range = discovery.get('date_range') or {}
    results = rules.get('results', [])
    hits = [r for r in results if r.get('hit') is True]
    doc = Document(); sec = doc.sections[0]
    sec.top_margin = Inches(.7); sec.bottom_margin = Inches(.7)
    doc.styles['Normal'].font.name = 'Arial'; doc.styles['Normal'].font.size = Pt(11)
    doc.add_heading('兴趣岛渠道经营分析报告', 0)
    doc.add_paragraph(f"观察期：{date_range.get('start', '未知')} 至 {date_range.get('end', '未知')}")
    doc.add_paragraph(f"数据来源：{discovery.get('source_file', '当前输入')}；Sheet：{discovery.get('source_sheet', '未知')}；数据行：{discovery.get('row_count', 0)}；列数：{len(discovery.get('headers', []))}")
    doc.add_heading('一、执行结果', level=1)
    if not results:
        doc.add_paragraph('本次输入没有满足规则所需的完整连续月份，Rules 未产生正式规则结果。报告不生成历史经营结论；R3、R4、R5、R6、R8、R9、R36、R37、R61 均因数据覆盖不足而不可执行。')
    else:
        doc.add_paragraph(f'本次 Rule Engine 产生 {len(results)} 条结果，其中命中 {len(hits)} 条。以下内容仅来自本次 Rule Result。')
        for row in hits:
            doc.add_heading(f"{row.get('rule_id', '未知规则')}｜{row.get('rule_name', '')}", level=2)
            period = row.get('period', {})
            doc.add_paragraph(f"观察期：{period.get('start', '')} 至 {period.get('end', '')}")
            for evidence in row.get('evidence') or []: doc.add_paragraph(str(evidence), style='List Bullet')
    doc.add_heading('二、数据边界', level=1)
    doc.add_paragraph('本报告不引用仓库历史报告、预置 semantic blocks 或其他 run 的 Rule Result。')
    args.output.parent.mkdir(parents=True, exist_ok=True); doc.save(args.output); print(args.output)

if __name__ == '__main__': main()
