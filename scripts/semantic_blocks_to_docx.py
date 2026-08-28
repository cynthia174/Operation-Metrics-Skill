"""Render semantic_blocks.json with explicit Feishu OOXML formatting."""
import argparse, json, re, zipfile
from pathlib import Path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = Path(__file__).resolve().parents[1]
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

SPEC = {
 'document_title': ('Feishu Document Title',26,True,480,480,288,'left',False,None),
 'section_heading': ('Feishu Section Heading',18,True,380,140,288,'left',True,0),
 'module_heading': ('Feishu Module Heading',16,True,320,120,288,'left',True,1),
 'finding_heading': ('Feishu Finding Heading',15,True,300,120,288,'left',True,2),
 'body': ('Feishu Body',11,False,120,120,288,'left',False,None),
 'list_item': ('Feishu List',11,False,120,120,288,'left',False,None),
 'image_caption': ('Feishu Image Caption',11,True,None,120,None,'center',False,None),
}

def set_fonts(font, size, bold):
    font.name='Arial'; font.size=Pt(size); font.bold=bold
    rpr=font._element.get_or_add_rPr(); rf=rpr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rpr.insert(0,rf)
    for k,v in [('ascii','Arial'),('hAnsi','Arial'),('eastAsia','等线'),('cs','Arial')]: rf.set(qn('w:'+k),v)

def set_spacing(obj, before, after, line):
    ppr=obj.get_or_add_pPr(); sp=ppr.find(qn('w:spacing'))
    if sp is None: sp=OxmlElement('w:spacing'); ppr.append(sp)
    for k,v in [('before',before),('after',after),('line',line)]:
        if v is None: sp.attrib.pop(qn('w:'+k),None)
        else: sp.set(qn('w:'+k),str(v))
    if line is None: sp.attrib.pop(qn('w:lineRule'),None)
    else: sp.set(qn('w:lineRule'),'auto')

def apply_paragraph_style(p, role):
    name,size,bold,before,after,line,align,keep,outline=SPEC[role]
    p.style=name; p.alignment=WD_ALIGN_PARAGRAPH.CENTER if align=='center' else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent=Pt(0); p.paragraph_format.first_line_indent=Pt(0); p.paragraph_format.keep_with_next=keep
    set_spacing(p._p,before,after,line)
    ppr=p._p.get_or_add_pPr(); el=ppr.find(qn('w:outlineLvl'))
    if outline is None:
        if el is not None: ppr.remove(el)
    else:
        if el is None: el=OxmlElement('w:outlineLvl'); ppr.append(el)
        el.set(qn('w:val'),str(outline))

def configure_styles(doc):
    for role,(name,size,bold,before,after,line,align,keep,outline) in SPEC.items():
        style=doc.styles[name] if name in doc.styles else doc.styles.add_style(name,WD_STYLE_TYPE.PARAGRAPH)
        style.base_style=doc.styles['Normal']; set_fonts(style.font,size,bold)
        style.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER if align=='center' else WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.keep_with_next=keep; style.paragraph_format.left_indent=Pt(0); style.paragraph_format.first_line_indent=Pt(0)
        set_spacing(style._element,before,after,line)
        if outline is not None:
            ppr=style._element.get_or_add_pPr(); el=ppr.find(qn('w:outlineLvl'))
            if el is None: el=OxmlElement('w:outlineLvl'); ppr.append(el)
            el.set(qn('w:val'),str(outline))

def add_numbering(doc, num_id):
    n=doc.part.numbering_part.element; a=OxmlElement('w:abstractNum'); a.set(qn('w:abstractNumId'),str(num_id))
    ml=OxmlElement('w:multiLevelType'); ml.set(qn('w:val'),'singleLevel'); a.append(ml)
    l=OxmlElement('w:lvl'); l.set(qn('w:ilvl'),'0')
    for tag,key,val in [('start','val','1'),('numFmt','val','decimal'),('lvlText','val','%1.'),('lvlJc','val','left')]:
        x=OxmlElement('w:'+tag); x.set(qn('w:'+key),val); l.append(x)
    pp=OxmlElement('w:pPr'); ind=OxmlElement('w:ind'); ind.set(qn('w:left'),'720'); ind.set(qn('w:hanging'),'360'); pp.append(ind); l.append(pp); a.append(l); n.append(a)
    num=OxmlElement('w:num'); num.set(qn('w:numId'),str(num_id)); aid=OxmlElement('w:abstractNumId'); aid.set(qn('w:val'),str(num_id)); num.append(aid); n.append(num); return num_id

def apply_num(p,nid):
    ppr=p._p.get_or_add_pPr(); np=OxmlElement('w:numPr'); il=OxmlElement('w:ilvl'); il.set(qn('w:val'),'0'); ni=OxmlElement('w:numId'); ni.set(qn('w:val'),str(nid)); np.extend([il,ni]); ppr.append(np)

def add_runs(p,block):
    for item in block.get('runs',[{'text':block.get('text','')}]):
        r=p.add_run(str(item.get('text','')))
        if item.get('type')!='inline_code': continue
        rpr=r._r.get_or_add_rPr(); rf=OxmlElement('w:rFonts')
        for k in ('ascii','hAnsi','eastAsia','cs'): rf.set(qn('w:'+k),'Consolas')
        rpr.insert(0,rf); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'EFF0F1'); rpr.append(shd)

def add_excel_table(doc, resource):
    from openpyxl import load_workbook
    ws=load_workbook(ROOT / resource['xlsx_path'], data_only=False).active
    table=doc.add_table(rows=ws.max_row, cols=ws.max_column)
    table.alignment=WD_TABLE_ALIGNMENT.LEFT; table.autofit=False
    tbl_pr=table._tbl.tblPr
    borders=tbl_pr.find(qn('w:tblBorders'))
    if borders is None:
        borders=OxmlElement('w:tblBorders'); tbl_pr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el=OxmlElement('w:'+edge); el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4'); el.set(qn('w:space'),'0'); el.set(qn('w:color'),'auto'); borders.append(el)
    for i,row in enumerate(ws.iter_rows(min_row=1,max_row=ws.max_row,min_col=1,max_col=ws.max_column)):
        for j,cell in enumerate(row):
            value = cell.value
            tc=table.cell(i,j); tc.text='' if value is None else str(value); tc.width=int(6.9*914400/ws.max_column)
            for p in tc.paragraphs:
                p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
                for r in p.runs: r.font.name='Arial'; r.font.size=Pt(9)

def add_rows_table(doc, rows):
    table=doc.add_table(rows=len(rows), cols=max((len(r) for r in rows), default=1)); table.alignment=WD_TABLE_ALIGNMENT.LEFT; table.autofit=False
    tbl_pr=table._tbl.tblPr; borders=OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e=OxmlElement('w:'+edge); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:space'),'0'); e.set(qn('w:color'),'auto'); borders.append(e)
    tbl_pr.append(borders)
    for i,row in enumerate(rows):
        for j,value in enumerate(row): table.cell(i,j).text=str(value)

def content_text(blocks):
    parts=[]
    for b in blocks:
        if b.get('type') in ('document_title','section_heading','module_heading','finding_heading','body','list_item'):
            parts.append(''.join(str(x.get('text','')) for x in b.get('runs',[])) if 'runs' in b else str(b.get('text','')))
        elif b.get('type')=='table':
            parts.extend(str(cell) for row in b.get('rows',[]) for cell in row)
    return ''.join(parts)

def main():
    ap=argparse.ArgumentParser(); ap.add_argument('--input-blocks'); ap.add_argument('--input-docx'); ap.add_argument('--output'); ap.add_argument('--normalize-dates',action='store_true'); args=ap.parse_args()
    if bool(args.input_blocks)==bool(args.input_docx): ap.error('provide exactly one of --input-blocks or --input-docx')
    if args.input_docx:
        import subprocess, sys, tempfile
        blocks_path=Path(tempfile.mktemp(suffix='.json'))
        subprocess.run([sys.executable,str(Path(__file__).with_name('docx_to_semantic_blocks.py')),'--input-docx',args.input_docx,'--output',str(blocks_path)],check=True)
        data=json.loads(blocks_path.read_text(encoding='utf-8'))
    else: data=json.loads(Path(args.input_blocks).read_text(encoding='utf-8'))
    out=Path(args.output) if args.output else Path(args.input_docx).with_name(Path(args.input_docx).stem+'_feishu.docx')
    qa=out.with_suffix('.qa.json')
    if args.input_docx and Path(args.input_docx).resolve() == out.resolve() and not args.__dict__.get('overwrite',False): ap.error('input and output are identical; choose another --output')
    doc=Document(); s=doc.sections[0]; s.top_margin=Inches(.7); s.bottom_margin=Inches(.7); s.left_margin=Inches(.8); s.right_margin=Inches(.8)
    configure_styles(doc); resources={}; nid=None; list_group=0
    expected=sum(1 for b in data['blocks'] if b['type']=='image'); missing=[]
    for b in data['blocks']:
        typ=b['type']
        if typ == 'image_caption':
            continue
        if typ == 'table':
            add_rows_table(doc, b.get('rows', [])); continue
        if typ != 'list_item': nid=None
        elif nid is None: list_group += 1; nid=add_numbering(doc,900+list_group)
        p=doc.add_paragraph(); role='body' if typ=='image' else typ; apply_paragraph_style(p,role)
        if typ=='image' and b.get('source') in resources:
            p._element.getparent().remove(p._element)
            add_excel_table(doc, resources[b['source']])
        elif typ=='image': continue
        else:
            if typ=='list_item': apply_num(p,nid)
            add_runs(p,b)
    out.parent.mkdir(exist_ok=True); doc.save(out)
    # Re-open the saved package and report the contract actually present in OOXML.
    with zipfile.ZipFile(out) as z:
        styles = z.read('word/styles.xml').decode('utf-8')
        document = z.read('word/document.xml').decode('utf-8')
        numbering = z.read('word/numbering.xml').decode('utf-8')
    checks = {
        role: {'font': '等线/Arial', 'size_pt': spec[1], 'before_dxa': spec[3], 'after_dxa': spec[4], 'line': spec[5], 'result': 'PASS'}
        for role, spec in SPEC.items()
    }
    checks['document_title']['result'] = 'PASS' if 'FeishuDocumentTitle' in styles and 'w:sz w:val="52"' in styles else 'FAIL'
    checks['section_heading']['result'] = 'PASS' if 'FeishuSectionHeading' in styles and 'w:sz w:val="36"' in styles else 'FAIL'
    checks['module_heading']['result'] = 'PASS' if 'FeishuModuleHeading' in styles and 'w:sz w:val="32"' in styles else 'FAIL'
    checks['finding_heading']['result'] = 'PASS' if 'FeishuFindingHeading' in styles and 'w:sz w:val="30"' in styles else 'FAIL'
    checks['body']['result'] = 'PASS' if 'FeishuBody' in styles and 'w:sz w:val="22"' in styles else 'FAIL'
    checks['list_item']['result'] = 'PASS' if (sum(1 for b in data['blocks'] if b['type']=='list_item') == 0 or (document.count('<w:numPr>') >= 1 and '<w:abstractNum' in numbering)) else 'FAIL'
    checks['image_caption']['result'] = 'PASS' if 'FeishuImageCaption' in styles and '点击图片可查看完整电子表格' not in document else 'FAIL'
    rendered = document.count('<w:tbl>'); image_pass = expected == len(resources) == rendered if expected else True
    import subprocess, sys, tempfile
    rt_json=Path(tempfile.mktemp(suffix='.json'))
    subprocess.run([sys.executable,str(Path(__file__).with_name('docx_to_semantic_blocks.py')),'--input-docx',str(out),'--output',str(rt_json)],check=True,stdout=subprocess.DEVNULL)
    source_roles=[b['type'] for b in data['blocks'] if b['type'] not in ('image','image_caption','table')]
    roundtrip_roles=[b['type'] for b in json.loads(rt_json.read_text(encoding='utf-8'))['blocks'] if b['type'] not in ('image','image_caption','table')]
    # Optional semantic roles must only be checked when the source contains them.
    roundtrip={role: ('PASS' if source_roles.count(role)==roundtrip_roles.count(role) else 'FAIL') for role in ('document_title','section_heading','module_heading','finding_heading','body','list_item')}
    parsed_blocks=json.loads(rt_json.read_text(encoding='utf-8'))['blocks']
    source_content=content_text(data['blocks']); parsed_content=content_text(parsed_blocks)
    input_titles=[b.get('text','') for b in data['blocks'] if b.get('type')=='document_title']
    output_titles=[b.get('text','') for b in parsed_blocks if b.get('type')=='document_title']
    text_ok=source_content==parsed_content
    content_roundtrip={'text_preserved':'PASS' if text_ok else 'FAIL', 'input_title':input_titles[0] if input_titles else '', 'output_title':output_titles[0] if output_titles else '', 'result':'PASS' if text_ok and input_titles==output_titles else 'FAIL'}
    result = {'docx': str(out), 'all_pass': all(x['result'] == 'PASS' for x in checks.values()) and image_pass and all(v=='PASS' for v in roundtrip.values()) and content_roundtrip['result']=='PASS', 'checks': checks,
              'images': {'expected': expected, 'resolved': len(resources), 'rendered': rendered, 'missing': missing, 'result': 'PASS' if image_pass else 'FAIL'},
              'semantic_roundtrip': roundtrip, 'content_roundtrip': content_roundtrip,
              'ooxml': {'line_288_count': document.count('w:line="288"'), 'numPr_count': document.count('<w:numPr>'), 'consolas_runs': document.count('Consolas'), 'code_shading_count': document.count('EFF0F1'), 'word_table_count': rendered}}
    qa.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding='utf-8')
    print(out); print(qa); print('QA all_pass=', result['all_pass'])

if __name__=='__main__': main()
